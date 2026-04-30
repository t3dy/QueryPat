"""DOCX -> markdown via python-docx.

Minimal converter: paragraphs by style, lists, headings. Tables become markdown tables.
Lossy for embedded images and complex formatting; sufficient for the archive.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document


def _para_to_md(para) -> str:
    text = para.text.strip()
    if not text:
        return ""
    style = (para.style.name or "").lower() if para.style else ""
    if style.startswith("heading 1") or style == "title":
        return f"# {text}"
    if style.startswith("heading 2"):
        return f"## {text}"
    if style.startswith("heading 3"):
        return f"### {text}"
    if style.startswith("heading 4"):
        return f"#### {text}"
    if style.startswith("heading"):
        return f"##### {text}"
    if "list" in style or "bullet" in style:
        return f"- {text}"
    return text


def _table_to_md(table) -> str:
    rows = []
    for row in table.rows:
        cells = [c.text.strip().replace("\n", " ").replace("|", "\\|") for c in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
    if not rows:
        return ""
    # Insert a header separator after the first row
    header_sep = "| " + " | ".join("---" for _ in table.rows[0].cells) + " |"
    return rows[0] + "\n" + header_sep + "\n" + "\n".join(rows[1:])


def convert(path: Path) -> tuple[str, str]:
    doc = Document(str(path))
    parts: list[str] = []
    # python-docx separates paragraphs and tables; iterate the body in order
    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}", 1)[-1]
        if tag == "p":
            for p in doc.paragraphs:
                if p._element is child:
                    md = _para_to_md(p)
                    if md:
                        parts.append(md)
                    break
        elif tag == "tbl":
            for t in doc.tables:
                if t._element is child:
                    md = _table_to_md(t)
                    if md:
                        parts.append(md)
                    break
    if not parts:
        raise ValueError("DOCX had no paragraphs or tables")
    return "\n\n".join(parts), "python-docx"
